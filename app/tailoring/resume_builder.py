from __future__ import annotations

import logging
import os
import re
import shutil
from pathlib import Path
from typing import TYPE_CHECKING

logger = logging.getLogger(__name__)

TAILORED_DIR = Path("resumes/tailored")


def extract_skill_categories(base_path: str) -> list[dict]:
    """Read the base DOCX and return the structured tech skills as a list of
    {category: str, skills: [str]} dicts, one per line in the TECHNICAL SKILLS section.

    Returns an empty list for PDF resumes (no structured parsing possible).
    """
    if Path(base_path).suffix.lower() != ".docx":
        return []
    import docx
    doc = docx.Document(base_path)
    return _parse_skill_categories(doc.paragraphs)


def _parse_skill_categories(paras) -> list[dict]:
    in_skills = False
    result: list[dict] = []
    for para in paras:
        text = para.text.strip()
        if not text:
            continue
        if text.upper() == "TECHNICAL SKILLS":
            in_skills = True
            continue
        if in_skills:
            if _is_section_header(text):
                break
            if ": " in text and para.runs and para.runs[0].bold:
                colon_idx = text.index(": ")
                category = text[:colon_idx]
                skills = [s.strip() for s in text[colon_idx + 2:].split(",") if s.strip()]
                result.append({"category": category, "skills": skills})
    return result


def build_tailored_resume(base_path: str, material: dict, job_id: str, resume_data: dict, job_title: str = "") -> str:
    """Generate a tailored resume file by injecting tailored content into the base resume.

    For DOCX: replaces the summary paragraph, experience bullets, and tech skills
    in-place, preserving all original formatting and hyperlinks (LinkedIn, GitHub, Portfolio).
    For PDF: cannot inject content — returns the original base path unchanged.

    The output filename follows the convention:
        {FirstName}_{LastName}_{YearsExp}_YEO_{Job_Title_Slug}.docx
    e.g. Deepak_Satuluri_3_YEO_Senior_Software_Engineer.docx

    Returns:
        Absolute path to the tailored resume file.
    """
    TAILORED_DIR.mkdir(parents=True, exist_ok=True)
    suffix = Path(base_path).suffix.lower()

    if suffix == ".docx":
        return _build_tailored_docx(base_path, material, job_id, resume_data, job_title)

    logger.warning(
        "Base resume is PDF — cannot inject tailored content. "
        "Convert to DOCX to enable per-application tailoring."
    )
    return base_path


# ── DOCX generation ──────────────────────────────────────────────────────────

def _build_tailored_docx(base_path: str, material: dict, job_id: str, resume_data: dict, job_title: str = "") -> str:
    import docx
    from app.config import get_settings

    settings = get_settings()
    first = settings.APPLICANT_FIRST_NAME or "Resume"
    last  = settings.APPLICANT_LAST_NAME  or ""
    years = settings.APPLICANT_YEARS_EXP  or "0"

    if job_title:
        # Normalize: "Senior Software Engineer - Backend" → "Senior_Software_Engineer"
        title_clean = re.split(r"\s*[-–|,]\s*", job_title.strip())[0].strip()
        title_slug  = re.sub(r"[^a-zA-Z0-9]+", "_", title_clean).strip("_")
        filename = f"{first}_{last}_{years}_YEO_{title_slug}.docx"
    else:
        filename = f"{job_id}.docx"

    out_path = str(TAILORED_DIR / filename)
    shutil.copy2(base_path, out_path)

    doc = docx.Document(out_path)
    paras = doc.paragraphs

    tailored_summary = material.get("summary", "").strip()
    tailored_bullets = [b.strip() for b in material.get("bullets", []) if b.strip()]
    tailored_tech_skills = material.get("tech_skills", [])

    # ── 1. Replace summary paragraph ─────────────────────────────────────────
    if tailored_summary:
        summary_idx = _find_summary_para_index(paras, resume_data.get("summary", ""))
        if summary_idx is not None:
            _replace_text(paras[summary_idx], tailored_summary)
            logger.info("Replaced summary at paragraph %d", summary_idx)
        else:
            logger.warning("Could not locate summary paragraph — skipping replacement")

    # ── 2. Replace experience bullets ─────────────────────────────────────────
    if tailored_bullets:
        bullet_indices = _find_experience_bullet_indices(paras)
        logger.info("Found %d bullet paragraphs; have %d tailored bullets", len(bullet_indices), len(tailored_bullets))

        for i, para_idx in enumerate(bullet_indices):
            if i < len(tailored_bullets):
                _replace_text(paras[para_idx], tailored_bullets[i])

        if len(tailored_bullets) > len(bullet_indices):
            logger.info(
                "%d tailored bullets generated but only %d slots in original — "
                "extra bullets dropped. Consider adding more bullets to base resume.",
                len(tailored_bullets), len(bullet_indices),
            )

    # ── 3. Replace tech skills lines ──────────────────────────────────────────
    if tailored_tech_skills:
        skill_indices = _find_tech_skill_para_indices(paras)
        logger.info("Found %d tech skill lines; have %d tailored categories", len(skill_indices), len(tailored_tech_skills))
        for i, para_idx in enumerate(skill_indices):
            if i < len(tailored_tech_skills):
                entry = tailored_tech_skills[i]
                _replace_skill_line(paras[para_idx], entry.get("category", ""), entry.get("skills", []))

    doc.save(out_path)
    logger.info("Saved tailored DOCX: %s", out_path)
    return out_path


def _replace_text(para, new_text: str) -> None:
    """Replace all paragraph content with new_text.

    Completely rebuilds the paragraph's inline content so that:
    - Hyperlink elements (e.g. quarkus/traefik in the original summary) are removed
      rather than left dangling after the replacement text.
    - The full replacement text is NOT accidentally made bold because the original
      paragraph's first run happened to be bold.

    For bullet paragraphs (those with <w:numPr>) the original style uses a bold
    action verb followed by normal body text, so we replicate that: the first word
    is written in a bold run and the remainder in a normal run.

    For all other paragraphs (summary, etc.) the entire text is written in a single
    normal-weight run.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    p = para._p
    pPr = p.find(qn('w:pPr'))
    is_bullet = pPr is not None and pPr.find(qn('w:numPr')) is not None

    # Collect rPr templates from existing runs before we wipe them
    normal_rPr = _pick_rPr(p, want_bold=False)
    bold_rPr = _pick_rPr(p, want_bold=True)

    # Remove every child that is not the paragraph properties element.
    # This includes <w:r> runs AND <w:hyperlink> elements so nothing bleeds through.
    for child in list(p):
        if child.tag != qn('w:pPr'):
            p.remove(child)

    if is_bullet and bold_rPr is not None:
        # Bold the first word (action verb); normal weight for the rest.
        parts = new_text.split(' ', 1)
        _append_run(p, parts[0], copy.deepcopy(bold_rPr))
        if len(parts) > 1:
            _append_run(p, ' ' + parts[1], copy.deepcopy(normal_rPr) if normal_rPr is not None else None)
    else:
        _append_run(p, new_text, copy.deepcopy(normal_rPr) if normal_rPr is not None else None)


def _append_run(p, text: str, rPr) -> None:
    """Append a <w:r> element with the given rPr and text to paragraph element p."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    r = OxmlElement('w:r')
    if rPr is not None:
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)


def _pick_rPr(p, want_bold: bool):
    """Return a copy of the first run's <w:rPr> that matches the bold preference.

    Iterates direct <w:r> children (skips hyperlinks) so we get a clean rPr
    representative of normal body text or bold text respectively.
    """
    from docx.oxml.ns import qn
    import copy

    for child in p:
        if child.tag != qn('w:r'):
            continue
        rPr = child.find(qn('w:rPr'))
        bold_el = rPr.find(qn('w:b')) if rPr is not None else None
        is_bold = bold_el is not None and bold_el.get(qn('w:val'), '1') not in ('0', 'false')
        if is_bold == want_bold:
            return copy.deepcopy(rPr) if rPr is not None else None
    return None


def _find_summary_para_index(paras, original_summary: str) -> int | None:
    """Find the index of the summary paragraph.

    Strategy:
    1. Find the 'PROFESSIONAL SUMMARY' section header.
    2. The very next non-empty paragraph is the summary body.
    3. Fallback: find the paragraph with highest word overlap to the original summary.
    """
    # Strategy 1 — look for the section header then take the next paragraph
    for i, para in enumerate(paras):
        text = para.text.strip()
        if text.upper() in ("PROFESSIONAL SUMMARY", "SUMMARY", "PROFILE", "OBJECTIVE"):
            for j in range(i + 1, min(i + 4, len(paras))):
                candidate = paras[j].text.strip()
                if len(candidate) > 40:
                    return j

    # Strategy 2 — word overlap with the original parsed summary
    if original_summary:
        orig_words = set(original_summary.lower().split())
        best_idx, best_score = None, 0.0
        for i, para in enumerate(paras[:40]):
            text = para.text.strip()
            if len(text) < 40 or _is_section_header(text) or _is_company_line(text):
                continue
            para_words = set(text.lower().split())
            denom = len(orig_words | para_words)
            score = len(orig_words & para_words) / denom if denom else 0
            if score > best_score:
                best_score, best_idx = score, i
        if best_score > 0.15:
            return best_idx

    return None


def _find_experience_bullet_indices(paras) -> list[int]:
    """Return indices of experience bullet paragraphs.

    Starts collecting after the 'PROFESSIONAL EXPERIENCE' header and stops
    at the next all-caps section header (e.g. 'KEY PROJECTS', 'EDUCATION').
    Skips job-title lines (contain tab or date patterns) and company lines (contain '•').
    """
    in_experience = False
    indices: list[int] = []

    for i, para in enumerate(paras):
        text = para.text.strip()
        if not text:
            continue

        if _is_section_header(text):
            if "EXPERIENCE" in text.upper():
                in_experience = True
            elif in_experience:
                # Hit the next section — stop collecting
                break
            continue

        if not in_experience:
            continue

        if _is_company_line(text) or _is_title_date_line(text):
            continue

        if len(text) > 40:
            indices.append(i)

    return indices


# ── Tech skills helpers ──────────────────────────────────────────────────────

def _find_tech_skill_para_indices(paras) -> list[int]:
    """Return indices of skill-line paragraphs inside the TECHNICAL SKILLS section.

    Each skill line has the pattern "Bold category: normal comma-separated values".
    Stops at the next all-caps section header.
    """
    in_skills = False
    indices: list[int] = []
    for i, para in enumerate(paras):
        text = para.text.strip()
        if not text:
            continue
        if text.upper() == "TECHNICAL SKILLS":
            in_skills = True
            continue
        if in_skills:
            if _is_section_header(text):
                break
            if ": " in text and para.runs and para.runs[0].bold:
                indices.append(i)
    return indices


def _replace_skill_line(para, category: str, skills: list[str]) -> None:
    """Replace a skill-line paragraph's content: bold 'Category: ' + normal 'skill1, skill2'."""
    from docx.oxml.ns import qn
    import copy

    p = para._p
    bold_rPr = _pick_rPr(p, want_bold=True)
    normal_rPr = _pick_rPr(p, want_bold=False)

    for child in list(p):
        if child.tag != qn('w:pPr'):
            p.remove(child)

    _append_run(p, f"{category}: ", copy.deepcopy(bold_rPr) if bold_rPr is not None else None)
    _append_run(p, ", ".join(skills), copy.deepcopy(normal_rPr) if normal_rPr is not None else None)


# ── Paragraph classifiers ────────────────────────────────────────────────────

def _is_section_header(text: str) -> bool:
    """All-caps short lines are section headers (PROFESSIONAL EXPERIENCE, etc.)."""
    t = text.strip()
    # Allow pipes and slashes in headers: "SENIOR BACKEND ENGINEER | DISTRIBUTED SYSTEMS"
    clean = t.replace("|", "").replace("&", "").replace("-", "").replace("/", "")
    return len(t) > 2 and clean.strip() == clean.strip().upper() and len(t.split()) <= 8


def _is_company_line(text: str) -> bool:
    """Company/location lines contain a bullet separator."""
    return "•" in text


def _is_title_date_line(text: str) -> bool:
    """Job title + date lines typically contain a tab or an en/em dash with a month."""
    months = {"jan","feb","mar","apr","may","jun","jul","aug","sep","oct","nov","dec"}
    lower = text.lower()
    has_month = any(m in lower for m in months)
    return "\t" in text or (has_month and ("–" in text or "-" in text or "present" in lower))
